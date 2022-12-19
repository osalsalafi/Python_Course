# ---------------------------------------------------------------------------------------------------
# -- Python Course => Handling of Word and PDF documents => 05. Create and write to Word documents --
# ---------------------------------------------------------------------------------------------------
import docx
mydoc = docx.Document()
mydoc.add_paragraph('This is Engineer Osama')

mydoc.save('WordFileTest.docx')

# if you want to add sentence to the same paragraph
second_paragraph = mydoc.add_paragraph('I am 26 years old')
second_paragraph.add_run(', But still young')

# if you want to add heading
mydoc.add_heading('Chapter 1',0)
mydoc.add_heading('Chapter 2',1)
mydoc.add_heading('Chapter 3',2)
mydoc.add_heading('Chapter 4',3)
mydoc.add_heading('Chapter 4',4)

print(mydoc.paragraphs[0].style) # _ParagraphStyle('Normal') id: 1958180647696
print(mydoc.paragraphs[3].style) # _ParagraphStyle('Heading 1') id: 1958180647696

# Change style
mydoc.paragraphs[0].style = mydoc.styles['Title']

# Delete style
# mydoc.paragraphs[2].style.delete()

mydoc.save('WordFileTest.docx')

# add picture
mydoc.add_picture('Osama.jpg', width=docx.shared.Inches(5), height=docx.shared.Inches(5))

mydoc.save('WordFileTest.docx')
