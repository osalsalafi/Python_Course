# --------------------------------------------------------------------------------------------------
# -- Python Course => Handling of Word and PDF documents => 04. Recognize and read Word documents --
# --------------------------------------------------------------------------------------------------

import docx

doc = docx.Document('academy_1.docx')

# TO KNOW HOW MANY PARAGRAPHS IN THE FILE
print(len(doc.paragraphs)) # 3

# TO PRINT THE FIRST PARAGRAPH
print(doc.paragraphs[0].text) # Academy 1

# TO PRINT THE RUNS OF THE PARAGRAPH
print(len(doc.paragraphs[1].runs)) # 1
print(doc.paragraphs[1].runs[0].text) #This is the file that I am using as a test for academy one

# print all the texts in the file
def get_text(doc):
    FileText = []
    for para in doc.paragraphs :
        FileText.append(para.text)
    return '\n'.join(FileText)

print(get_text(doc))
# # Academy 1
# This is the file that I am using as a test for academy one
# You can try to test the commands on this file to learn how to use the word file in python
